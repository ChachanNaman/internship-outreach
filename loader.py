"""Load the Talent Acquisition xlsx, derive domain/tier/agency, and store in SQLite."""

from __future__ import annotations

import logging
import re
from typing import Any, Optional
from urllib.parse import urlsplit
from zipfile import BadZipFile

import pandas as pd
from docx import Document

from db import Database, row_key

logger = logging.getLogger("outreach")

EXPECTED_COLUMNS = [
    "Name",
    "Job Title",
    "Linkedin URL",
    "Company Name",
    "Status",
    "Applied for Internship/Job",
    "Company Website",
    "Company Linkedin",
    "Company Social",
    "Company Twitter",
    "Location",
    "Company Niche",
]

TIER_1_KEYWORDS = ["pune"]
TIER_2_KEYWORDS = ["bengaluru", "bangalore"]
TIER_3_KEYWORDS = ["hyderabad", "mumbai", "delhi", "gurgaon", "gurugram", "noida", "ncr"]

AGENCY_NICHE = "staffing & recruiting"

# "Company Website" sometimes actually contains a social/platform profile URL
# (e.g. a LinkedIn company page) instead of the company's real site. Guessing
# emails at these domains produces meaningless/misdirected addresses (nobody's
# work email is @linkedin.com), so treat them as "no domain".
NON_COMPANY_DOMAINS = {
    "linkedin.com",
    "facebook.com",
    "twitter.com",
    "x.com",
    "instagram.com",
    "youtube.com",
    "github.com",
    "crunchbase.com",
    "medium.com",
    "wikipedia.org",
}


class XlsxLoadError(RuntimeError):
    """Raised when the source spreadsheet can't be read."""


def _truthy(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    if isinstance(value, (int, float)):
        return bool(value)
    text = str(value).strip().lower()
    return text in {"true", "yes", "y", "applied", "1"}


def extract_domain(website: Any) -> Optional[str]:
    if not website or (isinstance(website, float)):
        return None
    text = str(website).strip()
    if not text:
        return None
    if "://" not in text:
        text = "http://" + text
    try:
        netloc = urlsplit(text).netloc.lower()
    except ValueError:
        return None
    netloc = netloc.split("@")[-1]  # strip any userinfo
    netloc = netloc.split(":")[0]  # strip port
    if netloc.startswith("www."):
        netloc = netloc[4:]
    if not netloc or "." not in netloc:
        return None
    if netloc in NON_COMPANY_DOMAINS:
        return None
    return netloc


def classify_tier(location: Any) -> int:
    text = str(location or "").strip().lower()
    if any(k in text for k in TIER_1_KEYWORDS):
        return 1
    if any(k in text for k in TIER_2_KEYWORDS):
        return 2
    if any(k in text for k in TIER_3_KEYWORDS):
        return 3
    return 4


def is_agency(niche: Any) -> bool:
    return str(niche or "").strip().lower() == AGENCY_NICHE


def read_xlsx(path: str) -> pd.DataFrame:
    try:
        df = pd.read_excel(path, engine="openpyxl", dtype=str)
    except BadZipFile as exc:
        raise XlsxLoadError(
            f"'{path}' is not a valid .xlsx file (corrupt zip container). "
            "Re-export/re-save the spreadsheet from Excel/Google Sheets and try again."
        ) from exc
    except FileNotFoundError as exc:
        raise XlsxLoadError(f"Source spreadsheet not found: '{path}'") from exc
    except Exception as exc:  # noqa: BLE001 - surfaced as a clear CLI error
        raise XlsxLoadError(f"Could not read '{path}': {exc}") from exc

    df.columns = [str(c).strip() for c in df.columns]
    missing = [c for c in EXPECTED_COLUMNS if c not in df.columns]
    if missing:
        raise XlsxLoadError(
            f"'{path}' is missing expected column(s): {missing}. "
            f"Found columns: {list(df.columns)}"
        )
    return df


def load_and_prioritize(xlsx_path: str, db: Database) -> dict[str, int]:
    """Read the xlsx, classify every row, and insert new contacts into SQLite.

    Idempotent: rows already present (by row_key) are left untouched.
    """
    df = read_xlsx(xlsx_path)

    inserted = 0
    skipped_applied = 0
    already_known = 0

    for _, r in df.iterrows():
        name = (r.get("Name") or "").strip() if isinstance(r.get("Name"), str) else str(r.get("Name") or "").strip()
        company_name = str(r.get("Company Name") or "").strip()
        linkedin_url = str(r.get("Linkedin URL") or "").strip()

        if not name or not company_name:
            logger.warning("Skipping row with missing Name/Company Name")
            continue

        key = row_key(name, company_name, linkedin_url)
        applied = _truthy(r.get("Applied for Internship/Job"))
        website = r.get("Company Website")
        domain = extract_domain(website)
        niche = str(r.get("Company Niche") or "").strip()

        fields = {
            "row_key": key,
            "name": name,
            "job_title": str(r.get("Job Title") or "").strip(),
            "linkedin_url": linkedin_url,
            "company_name": company_name,
            "company_website": str(website or "").strip(),
            "company_domain": domain,
            "company_linkedin": str(r.get("Company Linkedin") or "").strip(),
            "company_social": str(r.get("Company Social") or "").strip(),
            "company_twitter": str(r.get("Company Twitter") or "").strip(),
            "location": str(r.get("Location") or "").strip(),
            "company_niche": niche,
            "source_status": str(r.get("Status") or "").strip(),
            "tier": classify_tier(r.get("Location")),
            "agency": int(is_agency(niche)),
            "applied_original": int(applied),
            "status": "skipped_applied" if applied else "new",
        }

        was_new = db.insert_contact_if_new(fields)
        if was_new:
            inserted += 1
            if applied:
                skipped_applied += 1
        else:
            already_known += 1

    logger.info(
        "Load complete: %d new contacts inserted (%d already-applied skipped), %d already known",
        inserted,
        skipped_applied,
        already_known,
    )
    return {
        "inserted": inserted,
        "skipped_applied": skipped_applied,
        "already_known": already_known,
        "total_rows": len(df),
    }


def load_docx_contacts(docx_path: str, db: Database) -> dict[str, int]:
    """Load a (Name, Email, Title, Company) table from a .docx into SQLite.

    Unlike the xlsx source, rows here already carry a real email address
    rather than needing name+domain pattern-guessing -- that known email is
    stored as the sole candidate so the verify step checks it directly
    instead of generating and testing guessed patterns.

    No Location/Company Niche data is available from this source, so rows
    default to tier 4 and get no focus-area sentence (dropped, same as any
    unmapped niche). Idempotent via the same row_key dedup as the xlsx loader.
    """
    try:
        document = Document(docx_path)
    except Exception as exc:  # noqa: BLE001
        raise XlsxLoadError(f"Could not read '{docx_path}': {exc}") from exc

    if not document.tables:
        raise XlsxLoadError(f"'{docx_path}' has no table to read contacts from.")

    table = document.tables[0]
    rows = [[c.text.strip() for c in r.cells] for r in table.rows]
    if not rows:
        raise XlsxLoadError(f"'{docx_path}' table is empty.")

    header = [h.strip().lower() for h in rows[0]]
    required = {"name", "email", "title", "company"}
    if not required.issubset(header):
        raise XlsxLoadError(
            f"'{docx_path}' table is missing expected column(s). "
            f"Found: {rows[0]}, need: {sorted(required)}"
        )
    idx = {col: header.index(col) for col in required}

    inserted = 0
    skipped_invalid = 0
    already_known = 0

    for r in rows[1:]:
        name = r[idx["name"]].strip()
        email = r[idx["email"]].strip()
        title = r[idx["title"]].strip()
        company_name = r[idx["company"]].strip()

        if not name or not company_name or "@" not in email:
            skipped_invalid += 1
            continue

        domain = email.split("@")[-1].strip().lower()
        key = row_key(name, company_name, email)

        fields = {
            "row_key": key,
            "name": name,
            "job_title": title,
            "linkedin_url": "",
            "company_name": company_name,
            "company_website": "",
            "company_domain": domain if domain in NON_COMPANY_DOMAINS else (domain or None),
            "company_linkedin": "",
            "company_social": "",
            "company_twitter": "",
            "location": "",
            "company_niche": "",
            "source_status": "",
            "source": "docx",
            "tier": 4,
            "agency": 0,
            "applied_original": 0,
            "status": "new",
            "candidate_emails": [email],
        }
        # Same social-platform-domain guard as the xlsx path -- a bogus
        # "email" whose domain is a social platform isn't a real company
        # mailbox either.
        if domain in NON_COMPANY_DOMAINS:
            fields["company_domain"] = None
            fields["candidate_emails"] = None

        was_new = db.insert_contact_if_new(fields)
        if was_new:
            inserted += 1
        else:
            already_known += 1

    logger.info(
        "Docx load complete: %d new contacts inserted, %d invalid rows skipped, %d already known",
        inserted,
        skipped_invalid,
        already_known,
    )
    return {
        "inserted": inserted,
        "skipped_invalid": skipped_invalid,
        "already_known": already_known,
        "total_rows": len(rows) - 1,
    }
